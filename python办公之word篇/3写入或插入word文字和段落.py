from docx import Document

document1 = Document(r'C:\Users\肖阿强\Desktop\练习.docx')

print('-----------1添加标题-----------')
document1.add_heading('3我是一级标题',level=1)      #添加一级标题


print('-----------2添加段落-----------')
document1.add_paragraph('我是正文')


print('-----------3添加分页符-----------')
document1.add_page_break()  #添加分页符

print('-----------4加块-----------')
zhw = document1.add_paragraph('后面这段文字含有格式：')
zhw.add_run('加粗').bold = True
zhw.add_run('普通')
zhw.add_run('斜体').italic = True
print(zhw.runs)                          #列表

print('-----------5段落插入-----------')
dl = document1.paragraphs[0]
dl.insert_paragraph_before('段前插入')





document1.save(r'C:\Users\肖阿强\Desktop\练习2.docx')  #创建并保存存到练习2