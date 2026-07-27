from docx import Document
from openpyxl import Workbook, load_workbook


# 获取特定列col在1-rows行区间的各单元格且值非None的行次信息
def list_rows(wb_exl, col, rows, value=None):
    wb = load_workbook(wb_exl)
    # 激活相应工作表附注
    ws = wb['附注']
    wb.active = wb.index(ws)
    list_r = []
    if not value:
        # 遍历col列1-rows行各单元格value值是否为None
        for row in range(1, rows + 1):
            if ws.cell(row, col).value:
                list_r.append(row)
        return list_r
    else:
        for row in range(1, rows + 1):
            if ws.cell(row, col).value == value:
                list_r.append(row)
        return list_r


# 获取第n个注释的表格数/第n个表格的行数
def table_num(list_rows_a, list_row_b, n):
    list_table = []
    if n > len(list_rows_a):
        print('超出表格个数')
    elif n == len(list_rows_a):
        for row in list_row_b:
            if row >= list_rows_a[n - 1]:
                list_table.append(row)
    else:
        for row in list_row_b:
            if row >= list_rows_a[n - 1] and row < list_rows_a[n]:
                list_table.append(row)
    return len(list_table)


# 查找特定段落后的第一个表格
def find_table_after_paragraph(doc, paragraph_index):
    target_paragraph = doc.paragraphs[paragraph_index]
    parent = target_paragraph._element.getparent()
    start_index = parent.index(target_paragraph._element) + 1

    for element in parent[start_index:]:
        if element.tag.endswith('tbl'):
            # 通过遍历文档的tables属性匹配表格
            for table in doc.tables:
                if table._element is element:
                    return table
    return None


doc = Document(r'C:\Users\肖阿强\Desktop\报告.docx')
wb = load_workbook(r'C:\Users\肖阿强\Desktop\附注.xlsx')
ws = wb['附注']
wb.active = wb.index(ws)
b = list_rows(r'C:\Users\肖阿强\Desktop\附注.xlsx',2,50)
c = list_rows(r'C:\Users\肖阿强\Desktop\附注.xlsx',3,50)
d = list_rows(r'C:\Users\肖阿强\Desktop\附注.xlsx',4,50)
d_title = list_rows(r'C:\Users\肖阿强\Desktop\附注.xlsx',4,50,'标题')

e = []
for i in range(len(b)):
    for j,t in enumerate(doc.paragraphs):
        if t.text == '应收账款':
            e.append(j)
print(e)

# ws.cell(2,b[i]).value




