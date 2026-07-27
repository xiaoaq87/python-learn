from docx import Document

doc = Document(r'C:\Users\肖阿强\Desktop\报告.docx')

for i,para in enumerate(doc.paragraphs):
    if para.text == '货币资金':
        print(i,para.text)

# for row in doc.tables[12].rows:
#     for cell in row.cells:
#         print(cell.text)




