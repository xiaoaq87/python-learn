
'''

paragraphs runs tables

文件.paragraphs 段落
段落.style.name 段落标题名字
段落.text 段落内容

'''


from docx import Document
import re
from docx.enum.style import WD_STYLE_TYPE    #用于3获取文件的结构

document1 = Document(r'C:\Users\肖阿强\Desktop\练习.docx')   #文件后缀记得加上

print('-----------1获取所有的标题-----------')
for para in document1.paragraphs:
    # if para.style.name == 'Heading 1':      #style 属性 标识标题
    #     print(para.text)                   #text ：内容

    if re.match(r'^Heading \d+$',para.style.name):
        print(para.text)

print('-----------2获取正文----------')
for para in document1.paragraphs:
    if para.style.name == 'Normal' and len(para.runs) >=1 :
        print(para.text)
        print(para.runs[1].text)


# print('-----------3获取文件的结构----------')

# style1 = document1.styles
# for i in style1:
#     if i.type == WD_STYLE_TYPE.PARAGRAPH:
#         print(i.name)

