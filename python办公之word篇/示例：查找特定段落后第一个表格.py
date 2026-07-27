from docx import Document


def find_table_after_paragraph(doc, paragraph_index):
    target_paragraph = doc.paragraphs[paragraph_index]
    parent = target_paragraph._element.getparent()
    start_index = parent.index(target_paragraph._element) + 1

    for element in parent[start_index:]:
        if element.tag.endswith('tbl'):
            # 通过遍历文档的tables属性匹配表格
            for table in doc.tables:
                if table._element is element:
                    return table
    return None


# 示例用法
doc = Document("example.docx")
table = find_table_after_paragraph(doc, 520)
if table:
    for row in table.rows:
        print([cell.text for cell in row.cells])