from docx import Document
from docx.shared import Cm

list1 = [
    ['姓名','性别','年龄'],
['孙兴华','男','20'],
['赵丽颖','女','23'],
['叶问','男','120'],
['李小龙','男','70'],
]
print('-----------1创建表并填入相应数据-----------')
document1 = Document(r'C:\Users\肖阿强\Desktop\练习.docx')
r=5
c=3
table1 = document1.add_table(rows=r,cols=c)
print(table1.rows)
for i in range(len(table1.rows)):
    cell1 = table1.rows[i].cells           #table1.rows是行对象  cell1是元组

    for j in range(len(table1.columns)):  #table1.columns不是table1.cols
        cell1[j].text = str(list1[i][j])  #注意要加str

print('-----------2增加行列-----------')

document1.tables[0].add_row()
document1.tables[0].add_column(Cm(5))
print(len(document1.tables[0].rows))

print('-----------3删除行-----------')
#删除第二行
print(len(document1.tables[0].rows),len(document1.tables[0].columns))
r2 = document1.tables[0].rows[1]
r2._element.getparent().remove(r2._element)   #删除第二行
print(len(document1.tables[0].rows),len(document1.tables[0].columns))

print('-----------4删除列-----------')
for cell in document1.tables[0].columns[1].cells:
    cell._element.getparent().remove(cell._element)

print(len(document1.tables[0].rows),len(document1.tables[0].columns))  #虽然删除了列，但len依然显示不变，因为自动补了

print('-----------5删除表-----------')
document1.tables[0]._element.getparent().remove(document1.tables[0]._element)



document1.save(r'C:\Users\肖阿强\Desktop\练习2.docx')